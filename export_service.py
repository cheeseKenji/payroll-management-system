import os
import csv
from datetime import datetime


class ExportService:

    def __init__(self):
        self.export_dir = "exports"
        if not os.path.exists(self.export_dir):
            os.makedirs(self.export_dir)

    def _timestamp(self):
        return datetime.now().strftime("%Y%m%d_%H%M%S")

    def _rows_to_list(self, rows):
        """Convert sqlite3.Row list to plain list of dicts."""
        return [dict(row) for row in rows]

    # ================= CSV EXPORT =================

    def export_csv(self, columns, rows, report_name="report"):
        filename = f"{report_name}_{self._timestamp()}.csv"
        filepath = os.path.join(self.export_dir, filename)

        data = self._rows_to_list(rows)

        with open(filepath, "w", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=columns)
            writer.writeheader()
            for row in data:
                writer.writerow({col: row.get(col, "") for col in columns})

        return filepath

    # ================= EXCEL EXPORT =================

    def export_excel(self, columns, rows, report_name="report"):
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

            filename = f"{report_name}_{self._timestamp()}.xlsx"
            filepath = os.path.join(self.export_dir, filename)

            data = self._rows_to_list(rows)

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = report_name

            # Header style
            header_font  = Font(bold=True, color="FFFFFF", size=11)
            header_fill  = PatternFill("solid", fgColor="2563EB")
            header_align = Alignment(horizontal="center", vertical="center")
            thin = Side(style="thin", color="CCCCCC")
            border = Border(left=thin, right=thin, top=thin, bottom=thin)

            # Write headers
            for col_idx, col in enumerate(columns, start=1):
                cell = ws.cell(row=1, column=col_idx, value=col.replace("_", " ").title())
                cell.font   = header_font
                cell.fill   = header_fill
                cell.alignment = header_align
                cell.border = border

            # Write data rows
            for row_idx, row in enumerate(data, start=2):
                fill = PatternFill("solid", fgColor="F0F4FF") if row_idx % 2 == 0 else PatternFill("solid", fgColor="FFFFFF")
                for col_idx, col in enumerate(columns, start=1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=row.get(col, ""))
                    cell.fill   = fill
                    cell.border = border
                    cell.alignment = Alignment(horizontal="center")

            # Auto column width
            for col in ws.columns:
                max_len = max((len(str(c.value)) if c.value else 0) for c in col)
                ws.column_dimensions[col[0].column_letter].width = max_len + 4

            # Summary row
            ws.append([])
            ws.append(["Report Generated:", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
            ws.append(["Total Records:", len(data)])

            wb.save(filepath)
            return filepath

        except ImportError:
            return self.export_csv(columns, rows, report_name)

    # ================= PDF EXPORT =================

    def export_pdf(self, columns, rows, report_name="report", title="Report"):
        try:
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.lib.styles import getSampleStyleSheet

            filename = f"{report_name}_{self._timestamp()}.pdf"
            filepath = os.path.join(self.export_dir, filename)

            data = self._rows_to_list(rows)

            doc = SimpleDocTemplate(filepath, pagesize=landscape(A4))
            styles = getSampleStyleSheet()
            elements = []

            # Title
            elements.append(Paragraph(f"<b>{title}</b>", styles["Title"]))
            elements.append(Paragraph(
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                styles["Normal"]
            ))
            elements.append(Spacer(1, 15))

            # Table data
            header = [col.replace("_", " ").title() for col in columns]
            table_data = [header]
            for row in data:
                table_data.append([str(row.get(col, "")) for col in columns])

            col_width = (landscape(A4)[0] - 60) / len(columns)
            col_widths = [col_width] * len(columns)

            t = Table(table_data, colWidths=col_widths, repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND",  (0, 0), (-1, 0),  colors.HexColor("#2563EB")),
                ("TEXTCOLOR",   (0, 0), (-1, 0),  colors.white),
                ("FONTNAME",    (0, 0), (-1, 0),  "Helvetica-Bold"),
                ("FONTSIZE",    (0, 0), (-1, 0),  9),
                ("ALIGN",       (0, 0), (-1, -1), "CENTER"),
                ("GRID",        (0, 0), (-1, -1), 0.4, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F0F4FF")]),
                ("FONTSIZE",    (0, 1), (-1, -1), 8),
            ]))

            elements.append(t)
            elements.append(Spacer(1, 20))
            elements.append(Paragraph(f"Total Records: {len(data)}", styles["Normal"]))

            doc.build(elements)
            return filepath

        except ImportError:
            return self.export_csv(columns, rows, report_name)

    # ================= WORD EXPORT =================

    def export_word(self, columns, rows, report_name="report", title="Report"):
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            filename = f"{report_name}_{self._timestamp()}.docx"
            filepath = os.path.join(self.export_dir, filename)

            data = self._rows_to_list(rows)

            doc = Document()

            # Title
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_para.add_run(title)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph()

            # Table
            table = doc.add_table(rows=1 + len(data), cols=len(columns))
            table.style = "Table Grid"

            # Header row
            hdr = table.rows[0]
            for i, col in enumerate(columns):
                cell = hdr.cells[i]
                cell.text = col.replace("_", " ").title()
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "2563EB")
                tcPr.append(shd)

            # Data rows
            for row_idx, row in enumerate(data):
                trow = table.rows[row_idx + 1]
                for col_idx, col in enumerate(columns):
                    trow.cells[col_idx].text = str(row.get(col, ""))

            doc.add_paragraph()
            doc.add_paragraph(f"Total Records: {len(data)}")

            doc.save(filepath)
            return filepath

        except ImportError:
            return self.export_csv(columns, rows, report_name)

    # ================= DB BACKUP =================

    def export_db_backup(self):
        import shutil
        from config import DATABASE_NAME

        filename = f"payroll_backup_{self._timestamp()}.db"
        filepath = os.path.join(self.export_dir, filename)
        shutil.copy2(DATABASE_NAME, filepath)
        return filepath
